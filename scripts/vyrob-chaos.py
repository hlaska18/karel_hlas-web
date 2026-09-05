#!/usr/bin/env python3
"""
Vyrobí složku CHAOS k pracovnímu listu 2 Digitální gramotnosti (Soubory a cloud).

CO TO JE. List posílá žáka: „Stáhni si složku CHAOS se soubory. Každý soubor
přesuň do správné složky podle obsahu a typu, přejmenuj ho, aby název říkal,
co je uvnitř, a drž se pravidel z teorie." Metodika dosud ukládala učiteli,
ať si těch 8–10 souborů vyrobí sám – jediná aktivita v obou tématech, kde
podklad nebyl hotový. Tohle ho dodává.

CO MUSÍ SOUBORY SPLŇOVAT, aby aktivita fungovala:

  * NÁZVY JSOU ŠPATNĚ SCHVÁLNĚ – přesně ty, které list zakazuje: „dokument1",
    „Nový textový dokument", název s mezerami a diakritikou, výchozí jméno
    z fotoaparátu, „Kopie – final2". Žák je má opravit.
  * OBSAH PROZRAZUJE PŘEDMĚT jednoznačně – jinak se nedá třídit. Každý soubor
    míří do jedné ze čtyř složek z listu: Informatika, Matematika, Cestina,
    Odborne_predmety.
  * RŮZNÉ TYPY – .docx, .txt, .csv, .jpg, .png, .pdf – protože list učí, že
    přípona říká, co je uvnitř.
  * VŠECHNO SMYŠLENÉ, ŽÁDNÁ OSOBA. Soubory vypadají jako žákovské, ale
    nenesou jméno, třídu ani nic, co by se dalo někomu přiřadit.

Obrázky se kreslí v PIL, dokumenty přes python-docx. Deterministicky – žádný
čas ani náhoda, aby se výstup v gitu neměnil, když se v něm nic nezměnilo.

PROČ ROVNOU ZIP A NE SLOŽKA S `_zdroje.txt`. Ta konvence existuje pro Grafiku,
kde pracovní listy odkazují na JEDNOTLIVÉ ukázky („Obrázky/foto_original.png"),
takže musí zůstat rozbalené na disku a jen se skryjí ze seznamu. Na CHAOS nikdo
po jednom neodkazuje – list posílá pro celou složku. Rozbalená kopie v `public/`
by tu tedy jen ležela navíc a v bance by dostala popisek „ukázky k porovnání",
který na soubory k roztřídění nesedí. Zdrojem pravdy je tenhle skript.

Spouštět z kořene projektu:
    python3 scripts/vyrob-chaos.py
"""

from __future__ import annotations

from pathlib import Path
import tempfile
import zipfile

from docx import Document
from docx.shared import Pt
from PIL import Image, ImageDraw, ImageFont

PODKLADY = Path("public/materialy/1L/2/Podklady k aktivitám")
ZIP = PODKLADY / "2. CHAOS.zip"

# Pevné datum, ať se ZIP v gitu nemění, když se obsah nezměnil – stejně jako
# ve `vyrob-ukazky-zip.py`.
CAS = (2026, 1, 1, 0, 0, 0)

# Font s českými znaky. macOS má tyhle; první, který existuje, vyhrává.
FONTY = [
    "/System/Library/Fonts/Supplemental/Arial.ttf",
    "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    "/Library/Fonts/Arial.ttf",
    "/System/Library/Fonts/Helvetica.ttc",
]


def font(velikost: int) -> ImageFont.FreeTypeFont:
    for f in FONTY:
        if Path(f).exists():
            return ImageFont.truetype(f, velikost)
    return ImageFont.load_default()


def docx_soubor(jmeno: str, nadpis: str, odstavce: list[str], tabulka: list[list[str]] | None = None) -> None:
    d = Document()
    d.add_heading(nadpis, level=1)
    for o in odstavce:
        p = d.add_paragraph(o)
        p.paragraph_format.space_after = Pt(6)
    if tabulka:
        t = d.add_table(rows=len(tabulka), cols=len(tabulka[0]))
        t.style = "Table Grid"
        for i, radek in enumerate(tabulka):
            for j, bunka in enumerate(radek):
                t.cell(i, j).text = bunka
    d.save(CIL / jmeno)


def txt_soubor(jmeno: str, text: str) -> None:
    (CIL / jmeno).write_text(text.strip() + "\n", encoding="utf-8")


# ── Kreslené soubory ────────────────────────────────────────────────────────

def obrazek_ozubene_kolo(jmeno: str) -> None:
    """Fotka výkresu ozubeného kola – strojírenství."""
    im = Image.new("RGB", (1200, 900), (246, 243, 236))
    d = ImageDraw.Draw(im)
    cx, cy, r = 520, 460, 240
    # rámeček výkresu a razítko
    d.rectangle([40, 40, 1160, 860], outline=(60, 60, 60), width=3)
    d.rectangle([760, 700, 1160, 860], outline=(60, 60, 60), width=2)
    d.text((780, 715), "OZUBENÉ KOLO", font=font(30), fill=(30, 30, 30))
    d.text((780, 760), "modul m = 2    z = 24", font=font(24), fill=(30, 30, 30))
    d.text((780, 795), "M 1:2    ČSN 01 4608", font=font(24), fill=(30, 30, 30))
    # zuby
    import math
    for i in range(24):
        a = 2 * math.pi * i / 24
        a2 = a + 2 * math.pi / 48
        body = [
            (cx + r * math.cos(a), cy + r * math.sin(a)),
            (cx + (r + 26) * math.cos(a + 0.05), cy + (r + 26) * math.sin(a + 0.05)),
            (cx + (r + 26) * math.cos(a2 - 0.05), cy + (r + 26) * math.sin(a2 - 0.05)),
            (cx + r * math.cos(a2), cy + r * math.sin(a2)),
        ]
        d.polygon(body, outline=(40, 40, 40), fill=(246, 243, 236))
    d.ellipse([cx - r, cy - r, cx + r, cy + r], outline=(40, 40, 40), width=3)
    d.ellipse([cx - 60, cy - 60, cx + 60, cy + 60], outline=(40, 40, 40), width=3)
    # osy (čerchovaně)
    for x in range(cx - r - 40, cx + r + 40, 24):
        d.line([x, cy, x + 14, cy], fill=(90, 90, 90), width=1)
    for y in range(cy - r - 40, cy + r + 40, 24):
        d.line([cx, y, cx, y + 14], fill=(90, 90, 90), width=1)
    # kóta
    d.line([cx - r - 26, cy + r + 70, cx + r + 26, cy + r + 70], fill=(40, 40, 40), width=2)
    d.text((cx - 40, cy + r + 78), "Ø 52", font=font(26), fill=(30, 30, 30))
    im.save(CIL / jmeno, quality=88)


def obrazek_krov(jmeno: str) -> None:
    """Fotka náčrtu krovu – stavebnictví."""
    im = Image.new("RGB", (1200, 900), (250, 248, 242))
    d = ImageDraw.Draw(im)
    # vaznice, krokve, hambalek
    d.line([150, 700, 1050, 700], fill=(50, 50, 50), width=6)          # vazný trám
    d.line([150, 700, 600, 220], fill=(50, 50, 50), width=6)           # krokev
    d.line([1050, 700, 600, 220], fill=(50, 50, 50), width=6)          # krokev
    d.line([375, 460, 825, 460], fill=(50, 50, 50), width=6)           # hambalek
    d.line([600, 220, 600, 700], fill=(120, 120, 120), width=2)        # osa
    d.text((640, 300), "krokev", font=font(28), fill=(40, 40, 40))
    d.text((560, 425), "hambalek", font=font(28), fill=(40, 40, 40))
    d.text((520, 715), "vazný trám", font=font(28), fill=(40, 40, 40))
    d.text((150, 100), "Krov – hambalková soustava, sklon 40°", font=font(34), fill=(30, 30, 30))
    d.text((150, 150), "náčrt do cvičení z pozemního stavitelství", font=font(24), fill=(80, 80, 80))
    im.save(CIL / jmeno, quality=86)


def obrazek_snimek_obrazovky(jmeno: str) -> None:
    """Snímek obrazovky Nastavení → Systém → O systému – informatika."""
    im = Image.new("RGB", (1280, 800), (243, 243, 243))
    d = ImageDraw.Draw(im)
    d.rectangle([0, 0, 1280, 40], fill=(255, 255, 255))
    d.text((20, 10), "Nastavení", font=font(18), fill=(40, 40, 40))
    d.rectangle([0, 40, 320, 800], fill=(235, 235, 235))
    for i, pol in enumerate(["Systém", "Bluetooth a zařízení", "Síť a internet", "Přizpůsobení", "Aplikace", "Účty"]):
        d.text((30, 70 + i * 44), pol, font=font(20), fill=(40, 40, 40) if i else (0, 90, 180))
    d.text((360, 60), "Systém  ›  O systému", font=font(30), fill=(30, 30, 30))
    d.text((360, 130), "Specifikace zařízení", font=font(22), fill=(30, 30, 30))
    radky = [
        ("Název zařízení", "UCEBNA-PC17"),
        ("Procesor", "Intel(R) Core(TM) i5-10400  2.90 GHz"),
        ("Nainstalovaná paměť RAM", "16,0 GB"),
        ("Typ systému", "64bitový operační systém"),
    ]
    for i, (k, v) in enumerate(radky):
        d.text((380, 180 + i * 40), k, font=font(19), fill=(90, 90, 90))
        d.text((720, 180 + i * 40), v, font=font(19), fill=(30, 30, 30))
    d.text((360, 380), "Specifikace Windows", font=font(22), fill=(30, 30, 30))
    for i, (k, v) in enumerate([("Edice", "Windows 11 Education"), ("Verze", "24H2")]):
        d.text((380, 430 + i * 40), k, font=font(19), fill=(90, 90, 90))
        d.text((720, 430 + i * 40), v, font=font(19), fill=(30, 30, 30))
    im.save(CIL / jmeno)


def pdf_sken(jmeno: str) -> None:
    """„Naskenovaná" stránka s příklady – matematika. PDF z obrázku."""
    im = Image.new("RGB", (1240, 1754), (252, 252, 250))
    d = ImageDraw.Draw(im)
    d.text((120, 120), "Lineární rovnice – procvičování", font=font(44), fill=(25, 25, 25))
    priklady = [
        "1)  3x + 7 = 22",
        "2)  5(x − 2) = 3x + 4",
        "3)  (x + 1)/2 = 4",
        "4)  2x − 9 = −x + 6",
        "5)  0,5x + 1,5 = 4",
    ]
    for i, p in enumerate(priklady):
        d.text((160, 260 + i * 140), p, font=font(38), fill=(25, 25, 25))
        d.line([160, 320 + i * 140, 1080, 320 + i * 140], fill=(200, 200, 200), width=1)
    d.text((120, 1080), "Zkouška: dosaď výsledek zpět do rovnice.", font=font(30), fill=(70, 70, 70))
    # lehké „naskenované" ztmavení okraje
    for k in range(18):
        d.rectangle([k, k, 1239 - k, 1753 - k], outline=(215 - k, 215 - k, 208 - k))
    im.save(CIL / jmeno, "PDF", resolution=150.0)


# ── Sestavení ───────────────────────────────────────────────────────────────

def main() -> None:
    global CIL
    PODKLADY.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory() as tmp:
        CIL = Path(tmp)
        sestav()
        soubory = sorted(CIL.iterdir(), key=lambda p: p.name)
        with zipfile.ZipFile(ZIP, "w", zipfile.ZIP_DEFLATED, compresslevel=9) as z:
            for s in soubory:
                info = zipfile.ZipInfo(f"CHAOS/{s.name}", CAS)
                info.compress_type = zipfile.ZIP_DEFLATED
                info.external_attr = 0o644 << 16
                z.writestr(info, s.read_bytes())
        for s in soubory:
            print(f"  {s.stat().st_size:7d} B  CHAOS/{s.name}")
    print(f"\n  → {ZIP}  ({ZIP.stat().st_size / 1024:.0f} kB)")


CIL = Path(".")


def sestav() -> None:

    # Čeština
    docx_soubor(
        "dokument1.docx",
        "Charakteristika literární postavy",
        [
            "Viktorka z Babičky Boženy Němcové je jednou z nejsilnějších postav české prózy "
            "19. století. Na začátku vyprávění je to veselá, hrdá dívka z Žernova, o kterou se "
            "ucházejí chlapci z celého okolí.",
            "Zlom přichází s příchodem černého myslivce. Němcová nepopisuje, co se přesně stalo – "
            "právě to nevyřčené dělá z Viktorky postavu, o které se dodnes píší úvahy.",
            "Vnější charakteristika: krásná, tmavovlasá, později zanedbaná. Vnitřní: hrdá, "
            "tajemná, po návratu uzavřená do vlastního světa u splavu.",
        ],
    )
    txt_soubor(
        "bez názvu.txt",
        """
        Erben – Kytice – Vodník (poznámky k rozboru)

        - balada: lyricko-epický žánr, tragický konec, nadpřirozená postava
        - motiv viny a trestu: matka varuje, dcera neposlechne
        - refrén a opakování zesilují napětí ("Nechoď, nechoď k jezeru")
        - kontrast: svět lidí (matka, chalupa) × svět vodníka (jezero, chlad)
        - závěr: rozhodnutí matky, tragédie dítěte – kdo je vinen?

        k dopsání: srovnat s Polednicí (také matka a dítě)
        """,
    )

    # Matematika
    txt_soubor(
        "Nový textový dokument.txt",
        """
        Kvadratická rovnice x^2 - 5x + 6 = 0

        a = 1, b = -5, c = 6
        D = b^2 - 4ac = 25 - 24 = 1
        x1 = (5 + 1) / 2 = 3
        x2 = (5 - 1) / 2 = 2

        zkouška: 9 - 15 + 6 = 0  OK
                 4 - 10 + 6 = 0  OK

        rozklad: (x - 2)(x - 3) = 0
        """,
    )
    txt_soubor(
        "tabulka (kopie).csv",
        """
        x;y
        -2;-1
        -1;1
        0;3
        1;5
        2;7
        3;9
        """,
    )
    pdf_sken("scan0001.pdf")

    # Informatika
    obrazek_snimek_obrazovky("snimek obrazovky 2026-03-17 v 9.41.33.png")
    docx_soubor(
        "úkol (1).docx",
        "Úkol z informatiky – příkazový řádek",
        [
            "Otevři příkazový řádek (Win + R, cmd) a přejdi do složky Documents.",
            "Vytvoř složku Mise a v ní tři podsložky: Vykresy, Dokumentace, Zalohy.",
            "Příkazem dir > obsah.txt ulož výpis do souboru a odevzdej ho spolu se snímkem obrazovky.",
            "Připomínka: del nemá koš. Co smažeš, je pryč.",
        ],
    )

    # Odborné předměty
    obrazek_ozubene_kolo("IMG_20260315_142211.jpg")
    obrazek_krov("DSC_0042.jpg")
    docx_soubor(
        "Kopie - final2.docx",
        "Protokol z měření – posuvné měřidlo",
        [
            "Úkol: změřit průměr a délku ocelového čepu posuvným měřidlem s přesností 0,05 mm.",
            "Měření se opakuje třikrát a zapisuje se aritmetický průměr.",
        ],
        tabulka=[
            ["měření", "průměr d [mm]", "délka l [mm]"],
            ["1", "19,95", "60,10"],
            ["2", "20,00", "60,05"],
            ["3", "19,95", "60,10"],
            ["průměr", "19,97", "60,08"],
        ],
    )


if __name__ == "__main__":
    main()
